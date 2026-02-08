from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import markdown2
from datetime import datetime
import io

def export_to_pdf(analysis_data, output_path=None):
    """
    Export analysis to PDF
    Returns: BytesIO buffer or saves to file
    """
    # Create buffer or file
    buffer = io.BytesIO() if not output_path else output_path
    
    # Create PDF
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for PDF elements
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#667eea',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor='#667eea',
        spaceAfter=12,
        spaceBefore=12
    )
    
    # Title
    title = Paragraph(analysis_data.get('title', 'Content Analysis Report'), title_style)
    story.append(title)
    story.append(Spacer(1, 0.3*inch))
    
    # Metadata
    method = Paragraph(f"<b>Analysis Method:</b> {analysis_data.get('method', 'N/A')}", styles['Normal'])
    story.append(method)
    story.append(Spacer(1, 0.1*inch))
    
    # Analysis statistics
    if 'analysis' in analysis_data:
        analysis = analysis_data['analysis']
        
        # Reading time
        if 'reading_time' in analysis:
            rt = analysis['reading_time']
            reading = Paragraph(
                f"<b>Reading Time:</b> {rt.get('reading_time', 'N/A')} ({rt.get('word_count', 0)} words)",
                styles['Normal']
            )
            story.append(reading)
            story.append(Spacer(1, 0.1*inch))
        
        # Sentiment
        if 'sentiment' in analysis:
            sent = analysis['sentiment']
            sentiment = Paragraph(
                f"<b>Sentiment:</b> {sent.get('sentiment', 'N/A')} ({sent.get('description', '')})",
                styles['Normal']
            )
            story.append(sentiment)
            story.append(Spacer(1, 0.1*inch))
        
        # Topics
        if 'topics' in analysis:
            topics_text = ", ".join([t.get('topic', '') for t in analysis['topics']])
            topics = Paragraph(f"<b>Topics:</b> {topics_text}", styles['Normal'])
            story.append(topics)
            story.append(Spacer(1, 0.1*inch))
    
    story.append(Spacer(1, 0.3*inch))
    
    # Executive Summary
    if analysis_data.get('executive_summary'):
        story.append(Paragraph("Executive Summary", heading_style))
        for item in analysis_data['executive_summary']:
            p = Paragraph(f"• {item}", styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 0.2*inch))
    
    # Detailed Summary
    if analysis_data.get('detailed_summary'):
        story.append(Paragraph("Detailed Summary", heading_style))
        for item in analysis_data['detailed_summary']:
            p = Paragraph(f"• {item}", styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.1*inch))
        story.append(Spacer(1, 0.2*inch))
    
    # Q&A Format
    if analysis_data.get('qa_format'):
        story.append(Paragraph("Questions & Answers", heading_style))
        for qa in analysis_data['qa_format']:
            q = Paragraph(f"<b>Q:</b> {qa.get('question', '')}", styles['Normal'])
            story.append(q)
            a = Paragraph(f"<b>A:</b> {qa.get('answer', '')}", styles['Normal'])
            story.append(a)
            story.append(Spacer(1, 0.15*inch))
    
    # Timeline Format
    if analysis_data.get('timeline'):
        story.append(Paragraph("Timeline", heading_style))
        for event in analysis_data['timeline']:
            timestamp = Paragraph(f"<b>{event.get('timestamp', '')}</b>", styles['Normal'])
            story.append(timestamp)
            desc = Paragraph(event.get('description', ''), styles['Normal'])
            story.append(desc)
            story.append(Spacer(1, 0.1*inch))
    
    # Insights Format
    if analysis_data.get('insights'):
        story.append(Paragraph("Key Insights", heading_style))
        for insight in analysis_data['insights']:
            p = Paragraph(f"💡 {insight}", styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.1*inch))
    
    # Confidence Score
    if analysis_data.get('confidence_score'):
        story.append(Spacer(1, 0.3*inch))
        confidence = Paragraph(
            f"<b>Confidence Score:</b> {analysis_data['confidence_score']}",
            styles['Normal']
        )
        story.append(confidence)
    
    # Footer
    story.append(Spacer(1, 0.5*inch))
    footer = Paragraph(
        f"<i>Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</i>",
        styles['Normal']
    )
    story.append(footer)
    
    # Build PDF
    doc.build(story)
    
    if not output_path:
        buffer.seek(0)
        return buffer
    
    return output_path

def export_to_docx(analysis_data, output_path=None):
    """
    Export analysis to Word document
    Returns: BytesIO buffer or saves to file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading(analysis_data.get('title', 'Content Analysis Report'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Analysis Method: {analysis_data.get('method', 'N/A')}")
    
    # Analysis statistics
    if 'analysis' in analysis_data:
        analysis = analysis_data['analysis']
        
        if 'reading_time' in analysis:
            rt = analysis['reading_time']
            doc.add_paragraph(
                f"Reading Time: {rt.get('reading_time', 'N/A')} ({rt.get('word_count', 0)} words)"
            )
        
        if 'sentiment' in analysis:
            sent = analysis['sentiment']
            doc.add_paragraph(
                f"Sentiment: {sent.get('sentiment', 'N/A')} - {sent.get('description', '')}"
            )
        
        if 'topics' in analysis:
            topics_text = ", ".join([t.get('topic', '') for t in analysis['topics']])
            doc.add_paragraph(f"Topics: {topics_text}")
    
    doc.add_paragraph()  # Space
    
    # Executive Summary
    if analysis_data.get('executive_summary'):
        doc.add_heading('Executive Summary', level=1)
        for item in analysis_data['executive_summary']:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()
    
    # Detailed Summary
    if analysis_data.get('detailed_summary'):
        doc.add_heading('Detailed Summary', level=1)
        for item in analysis_data['detailed_summary']:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()
    
    # Q&A Format
    if analysis_data.get('qa_format'):
        doc.add_heading('Questions & Answers', level=1)
        for qa in analysis_data['qa_format']:
            p = doc.add_paragraph()
            p.add_run(f"Q: {qa.get('question', '')}").bold = True
            doc.add_paragraph(f"A: {qa.get('answer', '')}")
            doc.add_paragraph()
    
    # Timeline Format
    if analysis_data.get('timeline'):
        doc.add_heading('Timeline', level=1)
        for event in analysis_data['timeline']:
            p = doc.add_paragraph()
            p.add_run(event.get('timestamp', '')).bold = True
            doc.add_paragraph(event.get('description', ''))
    
    # Insights Format
    if analysis_data.get('insights'):
        doc.add_heading('Key Insights', level=1)
        for insight in analysis_data['insights']:
            doc.add_paragraph(f"💡 {insight}", style='List Bullet')
    
    # Confidence Score
    if analysis_data.get('confidence_score'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"Confidence Score: {analysis_data['confidence_score']}").bold = True
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    )
    footer.runs[0].italic = True
    
    # Save to buffer or file
    if not output_path:
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    doc.save(output_path)
    return output_path

def export_to_markdown(analysis_data):
    """
    Export analysis to Markdown format
    Returns: string
    """
    md = []
    
    # Title
    md.append(f"# {analysis_data.get('title', 'Content Analysis Report')}\n")
    
    # Metadata
    md.append(f"**Analysis Method:** {analysis_data.get('method', 'N/A')}\n")
    
    # Analysis statistics
    if 'analysis' in analysis_data:
        analysis = analysis_data['analysis']
        
        if 'reading_time' in analysis:
            rt = analysis['reading_time']
            md.append(f"**Reading Time:** {rt.get('reading_time', 'N/A')} ({rt.get('word_count', 0)} words)\n")
        
        if 'sentiment' in analysis:
            sent = analysis['sentiment']
            md.append(f"**Sentiment:** {sent.get('sentiment', 'N/A')} - {sent.get('description', '')}\n")
        
        if 'topics' in analysis:
            topics_text = ", ".join([t.get('topic', '') for t in analysis['topics']])
            md.append(f"**Topics:** {topics_text}\n")
    
    md.append("\n---\n\n")
    
    # Executive Summary
    if analysis_data.get('executive_summary'):
        md.append("## Executive Summary\n\n")
        for item in analysis_data['executive_summary']:
            md.append(f"- {item}\n")
        md.append("\n")
    
    # Detailed Summary
    if analysis_data.get('detailed_summary'):
        md.append("## Detailed Summary\n\n")
        for item in analysis_data['detailed_summary']:
            md.append(f"- {item}\n")
        md.append("\n")
    
    # Q&A Format
    if analysis_data.get('qa_format'):
        md.append("## Questions & Answers\n\n")
        for qa in analysis_data['qa_format']:
            md.append(f"**Q:** {qa.get('question', '')}\n\n")
            md.append(f"**A:** {qa.get('answer', '')}\n\n")
    
    # Timeline Format
    if analysis_data.get('timeline'):
        md.append("## Timeline\n\n")
        for event in analysis_data['timeline']:
            md.append(f"**{event.get('timestamp', '')}**\n\n")
            md.append(f"{event.get('description', '')}\n\n")
    
    # Insights Format
    if analysis_data.get('insights'):
        md.append("## Key Insights\n\n")
        for insight in analysis_data['insights']:
            md.append(f"💡 {insight}\n\n")
    
    # Confidence Score
    if analysis_data.get('confidence_score'):
        md.append(f"\n**Confidence Score:** {analysis_data['confidence_score']}\n")
    
    # Footer
    md.append(f"\n---\n\n*Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}*\n")
    
    return "".join(md)

def export_to_json(analysis_data):
    """
    Export analysis to JSON format
    Returns: JSON string
    """
    export_data = {
        "title": analysis_data.get('title', ''),
        "method": analysis_data.get('method', ''),
        "generated_at": datetime.now().isoformat(),
        "analysis": analysis_data.get('analysis', {}),
        "executive_summary": analysis_data.get('executive_summary', []),
        "detailed_summary": analysis_data.get('detailed_summary', []),
        "qa_format": analysis_data.get('qa_format', []),
        "timeline": analysis_data.get('timeline', []),
        "insights": analysis_data.get('insights', []),
        "confidence_score": analysis_data.get('confidence_score', ''),
        "doc_id": analysis_data.get('doc_id', '')
    }
    
    return json.dumps(export_data, indent=2)