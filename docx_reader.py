from docx import Document
import re

def extract_text_from_docx(file_path):
    """
    Extract text from DOCX file
    Returns: (text, metadata)
    """
    try:
        doc = Document(file_path)
        
        # Extract metadata
        metadata = {
            "title": doc.core_properties.title or "Untitled Document",
            "author": doc.core_properties.author or "Unknown",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "Unknown",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "Unknown",
            "paragraphs": len(doc.paragraphs),
            "sections": len(doc.sections)
        }
        
        # Extract text from paragraphs
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        
        full_text = " ".join(text_parts)
        
        # Clean up whitespace
        full_text = re.sub(r'\s+', ' ', full_text).strip()
        
        if not full_text or len(full_text) < 50:
            raise Exception("No readable text found in DOCX file")
        
        return full_text, metadata
        
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")